from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mark28pro/ADS")
SOURCE = ROOT / "outputs/fase2/documento_fase_2.md"
OUT = ROOT / "outputs/fase2/Sistema_Inventario_Taller_San_Jose_Fase_2.docx"
SCREENSHOTS = ROOT / "outputs/fase2/screenshots"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_run_format(run, bold=False, italic=False, size=12, color="000000", font="Arial"):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_inline_markdown(paragraph, text, size=12):
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_format(run, bold=True, size=size)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_format(run, italic=True, size=size)
        else:
            run = paragraph.add_run(part)
            set_run_format(run, size=size)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after, color in [
        ("Title", 14, 0, 8, "000000"),
        ("Heading 1", 14, 18, 8, "000000"),
        ("Heading 2", 13, 14, 6, "000000"),
        ("Heading 3", 12, 10, 4, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc):
    title_lines = [
        "UNIVERSIDAD DON BOSCO",
        "ANÁLISIS Y DISEÑO DE SISTEMAS",
        "Docente: Juan Antonio Miranda Figueroa",
        "",
        "SISTEMA DE GESTIÓN DE INVENTARIO DE REPUESTOS",
        "TALLER MECÁNICO SAN JOSÉ",
    ]
    for i, line in enumerate(title_lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not line:
            p.add_run("")
            continue
        run = p.add_run(line)
        set_run_format(run, bold=i in (0, 1, 4, 5), size=14 if i in (0, 4, 5) else 12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTEGRANTE")
    set_run_format(run, bold=True, size=12)

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    set_cell_margins(table)
    headers = ["APELLIDOS", "NOMBRES", "CARNÉ"]
    values = ["García Villeda", "Enrique Alejandro", "GV200136"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "F1F3F4")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cell.paragraphs[0].add_run(text)
        set_run_format(r, bold=True, size=11)
    for idx, text in enumerate(values):
        cell = table.cell(1, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cell.paragraphs[0].add_run(text)
        set_run_format(r, size=11)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ciudadela Don Bosco, 21 de mayo de 2026.")
    set_run_format(run, size=12)
    doc.add_page_break()


def add_table_from_markdown(doc, lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(set(c) <= set("-: ") for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    set_cell_margins(table)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, "F1F3F4")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 and len(text) < 15 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markdown(p, text, size=10 if len(row) > 3 else 11)
            for run in p.runs:
                run.bold = r_idx == 0 or run.bold
    doc.add_paragraph()


def add_code_block(doc, code, language):
    label = "Diagrama en notación Mermaid" if language == "mermaid" else "Bloque técnico"
    p = doc.add_paragraph()
    r = p.add_run(label)
    set_run_format(r, bold=True, size=11, color="434343")
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run_format(r, size=9, font="Courier New", color="333333")


def add_screenshots(doc):
    doc.add_heading("12.3 Evidencia del sistema implementado", level=2)
    p = doc.add_paragraph()
    add_inline_markdown(
        p,
        "Las siguientes capturas corresponden a la versión funcional implementada con PHP, MySQL y Bootstrap. Se incluyen como evidencia del avance operativo usado para la defensa de la Fase 2.",
    )
    shots = [
        ("01-login.png", "Inicio de sesión"),
        ("02-dashboard.png", "Dashboard operativo"),
        ("03-repuestos.png", "Listado de repuestos"),
        ("04-movimientos.png", "Registro de movimientos"),
        ("05-proveedores.png", "Gestión de proveedores"),
        ("06-compras.png", "Listado de compras"),
        ("07-compra-form.png", "Formulario de compra"),
        ("08-clientes.png", "Gestión de clientes"),
        ("09-reportes.png", "Reportes"),
    ]
    for file_name, caption in shots:
        path = SCREENSHOTS / file_name
        if not path.exists():
            continue
        doc.add_heading(caption, level=2)
        doc.add_picture(str(path), width=Inches(6.35))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figura: {caption}.")
        set_run_format(run, italic=True, size=10, color="555555")


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lang = ""
    code_lines = []
    pending_table = []
    skip_cover = True
    screenshots_added = False

    for raw in lines:
        line = raw.rstrip()
        if skip_cover:
            if line.strip() == "# Índice":
                skip_cover = False
            else:
                continue

        if line.strip() == "\\pagebreak":
            if pending_table:
                add_table_from_markdown(doc, pending_table)
                pending_table = []
            doc.add_page_break()
            continue

        if line.startswith("```"):
            if not in_code:
                if pending_table:
                    add_table_from_markdown(doc, pending_table)
                    pending_table = []
                in_code = True
                code_lang = line.strip("`").strip()
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines), code_lang)
                in_code = False
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            pending_table.append(line)
            continue
        if pending_table:
            add_table_from_markdown(doc, pending_table)
            pending_table = []

        if not line.strip():
            continue

        if "Captura pendiente" in line or "reemplazar cada placeholder" in line:
            continue

        if line.startswith("# "):
            if line.strip() == "# 13. Referencias bibliográficas" and not screenshots_added:
                add_screenshots(doc)
                screenshots_added = True
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", line), size=12)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, line[2:], size=12)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_inline_markdown(p, line[2:], size=11)
        else:
            p = doc.add_paragraph()
            add_inline_markdown(p, line, size=12)

    if pending_table:
        add_table_from_markdown(doc, pending_table)

    if not screenshots_added:
        add_screenshots(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
